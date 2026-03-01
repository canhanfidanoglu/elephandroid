import io

from docx import Document
from docx.shared import Pt, Inches

from .models import PlanReport


def build_docx(report: PlanReport) -> bytes:
    """Build a DOCX document from a PlanReport and return as bytes."""
    doc = Document()

    # Title
    doc.add_heading(f"{report.plan_name} - Progress Report", level=0)

    # Subtitle
    doc.add_paragraph(f"Generated: {report.generated_at:%Y-%m-%d %H:%M}")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)

    not_started = report.total_tasks - report.completed_tasks - sum(
        b.in_progress for b in report.buckets
    )
    in_progress = sum(b.in_progress for b in report.buckets)

    doc.add_paragraph(
        f"This report covers {report.total_tasks} tasks across "
        f"{len(report.buckets)} buckets. "
        f"Overall completion is {report.overall_percentage:.0f}%. "
        f"{report.completed_tasks} tasks are completed, "
        f"{in_progress} are in progress, and "
        f"{not_started} have not started."
    )

    # Epic Progress
    if report.epics:
        doc.add_heading("Epic Progress", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Epic"
        hdr[1].text = "Total"
        hdr[2].text = "Completed"
        hdr[3].text = "%"
        for epic in report.epics:
            row = table.add_row().cells
            row[0].text = epic.name
            row[1].text = str(epic.total)
            row[2].text = str(epic.completed)
            row[3].text = f"{epic.percentage:.0f}%"

    # Bucket Breakdown
    if report.buckets:
        doc.add_heading("Bucket Breakdown", level=1)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Bucket"
        hdr[1].text = "Total"
        hdr[2].text = "Completed"
        hdr[3].text = "In Progress"
        hdr[4].text = "Not Started"
        for bucket in report.buckets:
            row = table.add_row().cells
            row[0].text = bucket.name
            row[1].text = str(bucket.total)
            row[2].text = str(bucket.completed)
            row[3].text = str(bucket.in_progress)
            row[4].text = str(bucket.not_started)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
